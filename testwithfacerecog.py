import pandas as pd
from docx import Document
from docx.shared import Inches , Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT 
from pillow_heif import register_heif_opener
import os
import requests
from io import BytesIO
from PIL import Image
import cv2
import numpy as np

register_heif_opener()

# Load the pre-trained face cascade classifier
face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')

def detect_and_crop_face(img, desired_aspect_ratio):
    # Convert PIL Image to OpenCV format
    if isinstance(img, Image.Image):
        img_cv = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)
    else:
        img_cv = img

    # Convert to grayscale for face detection
    gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
    
    # Detect faces
    faces = face_cascade.detectMultiScale(gray, 1.1, 4)
    
    if len(faces) > 0:
        # Get the largest face (assuming it's the main subject)
        largest_face = max(faces, key=lambda x: x[2] * x[3])
        x, y, w, h = largest_face
        
        # Calculate face center
        face_center_x = x + w//2
        face_center_y = y + h//2
        
        # Calculate crop dimensions based on desired aspect ratio
        img_height, img_width = img_cv.shape[:2]
        
        if img_width/img_height > desired_aspect_ratio:
            new_width = int(img_height * desired_aspect_ratio)
            # Center the crop window on the face
            left = max(0, min(face_center_x - new_width//2, img_width - new_width))
            crop_box = (left, 0, left + new_width, img_height)
        else:
            new_height = int(img_width / desired_aspect_ratio)
            # Center the crop window on the face
            top = max(0, min(face_center_y - new_height//2, img_height - new_height))
            crop_box = (0, top, img_width, top + new_height)
            
        return crop_box
    else:
        # Fall back to center-based cropping if no face detected
        img_aspect_ratio = img_width / img_height
        if img_aspect_ratio > desired_aspect_ratio:
            new_width = int(img_height * desired_aspect_ratio)
            left_margin = (img_width - new_width) // 2
            right_margin = img_width - new_width - left_margin
            crop_box = (left_margin, 0, img_width - right_margin, img_height)
        else:
            new_height = int(img_width / desired_aspect_ratio)
            top_margin = (img_height - new_height) // 4
            bottom_margin = img_height - new_height - top_margin*3
            crop_box = (0, top_margin, img_width, img_height - bottom_margin)
        return crop_box

df = pd.read_excel('test.xlsx')  
# df=pd.read_ex('fix.csv')
doc = Document()

# print(df.columns)

table = doc.add_table(rows=0, cols=3)
for row in table.rows:
    row.height_rule = Inches(5) 

# Fix pandas FutureWarnings by using proper assignment
quote_col = 'Submit a clean, creative yearbook quote (under 100 characters) to be printed under your image.'
photo_col = 'Upload a clear, well-lit, decent photo (1:1 ratio or passport size). Editing is not allowed, and you can only upload once.'

# Use proper assignment instead of inplace=True
df[quote_col] = df[quote_col].fillna('Lorem ipsum dolor sit amet, consectetur adipiscing elit')
df[photo_col] = df[photo_col].fillna('https://drive.google.com/file/d/1dxSlFP84x_jJbAH2SJMjOrwW_S87b6iy/view?usp=sharing')

# Check if Default.jpg exists, if not create a simple one
if not os.path.exists('Default.jpg'):
    print("Default.jpg not found. Creating a simple placeholder image...")
    # Create a simple white image
    img = Image.new('RGB', (300, 400), color='white')
    img.save('Default.jpg')
    print("Default.jpg created successfully.")

ratio = 1
# quotes = 'Lorem ipsum dolor sit amet, consectetur adipiscing elit'
for i in range(0, len(df), 3):
    row_cells = table.add_row().cells
    for j in range(3):
        if i + j < len(df):
            # print(i,j)
            photo_link = df.loc[i + j, photo_col]
            name = df.loc[i + j, 'Name']
            surname = df.loc[i + j, 'BITS ID (this form is only for students enrolled in the year 2022)']

                
            quotes = df.loc[i + j, quote_col]

            
            
        
            paragraph = row_cells[j].add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run()
            if(photo_link == 'https://drive.google.com/file/d/1dxSlFP84x_jJbAH2SJMjOrwW_S87b6iy/view?usp=sharing'):
                try:
                    image_data = cv2.imread('Default.jpg')
                    img = Image.fromarray(image_data)   
                    
                    desired_aspect_ratio = 1/ratio
                    img_aspect_ratio = img.width / img.height

                    if img_aspect_ratio > desired_aspect_ratio:
                        new_width = int(img.height * desired_aspect_ratio)
                        left_margin = (img.width - new_width) // 2
                        right_margin = img.width - new_width - left_margin
                        crop_box = (left_margin, 0, img.width - right_margin, img.height)
                    else:
                        new_height = int(img.width / desired_aspect_ratio)
                        top_margin = (img.height - new_height) // 4
                        bottom_margin = img.height - new_height - top_margin*3
                        crop_box = (0, top_margin, img.width, img.height - bottom_margin)

                    cropped_img = img.crop(crop_box)
                    new_width = img.width
                    new_height = int(new_width * ratio)

                    img = cropped_img.resize((new_width, new_height))

                    img_io = BytesIO()
                    img.save(img_io, format='JPEG')
                    img_io.seek(0) 
                    run.add_picture(img_io, width=Inches(1.5))            
                    paragraph.add_run('\n')
                    paragraph.add_run(f'{surname}\n').italic = True
                    paragraph.add_run(f'{name}\n').bold = True
                    quotes_run = paragraph.add_run(quotes)
                    quotes_run.font.size = Pt(10)
                except Exception as e:
                    print(f"Error processing default image: {e}")
                    # Add text only if image processing fails
                    paragraph.add_run(f'{surname}\n').italic = True
                    paragraph.add_run(f'{name}\n').bold = True
                    quotes_run = paragraph.add_run(quotes)
                    quotes_run.font.size = Pt(10)
                continue
            import requests
            from io import BytesIO
            
            file_id = ''
            try:
                # Modified to handle both 'open' and 'file' formats of Google Drive URLs
                if 'open?id=' in photo_link:
                    file_id = photo_link.split('open?id=')[1]
                else:
                    file_id = photo_link.split('=')[1]
                    
                download_url = f'https://drive.google.com/uc?export=download&id={file_id}'
                
                response = requests.get(download_url)
                image_data = BytesIO(response.content)
                from PIL import Image
                img = Image.open(image_data)
                
                desired_aspect_ratio = 1/ratio
                
                # Use face detection to determine crop box
                crop_box = detect_and_crop_face(img, desired_aspect_ratio)
                cropped_img = img.crop(crop_box)
                
                new_width = img.width
                new_height = int(new_width * ratio)

                img = cropped_img.resize((new_width, new_height))
                img = img.convert('RGB')
                # print(img)
                img_io = BytesIO()
                try:
                    img.save(img_io,format='JPEG')
                except:
                    try:
                        img.save(img_io,format='PNG')
                    except:
                        img.save(img_io,format='HEIC')
                img_io.seek(0)
                run.add_picture(img_io, width=Inches(1.5))            
                paragraph.add_run('\n')
                paragraph.add_run(f'{surname}\n').italic = True
                paragraph.add_run(f'{name}\n').bold = True
                quotes_run = paragraph.add_run(quotes)
                quotes_run.font.size = Pt(10)
            except Exception as e:   
                print(f"Error processing image for {name}: {e}")
                try:
                    # Try to use Default.jpg as fallback
                    run.add_picture('Default.jpg', width=Inches(1.5))            
                    paragraph.add_run('\n')
                    paragraph.add_run(f'{surname}\n').italic = True
                    paragraph.add_run(f'{name}\n').bold = True
                    quotes_run = paragraph.add_run(quotes)
                    quotes_run.font.size = Pt(10)
                except Exception as e2:
                    print(f"Error with Default.jpg: {e2}")
                    # Add text only if both image processing attempts fail
                    paragraph.add_run(f'{surname}\n').italic = True
                    paragraph.add_run(f'{name}\n').bold = True
                    quotes_run = paragraph.add_run(quotes)
                    quotes_run.font.size = Pt(10)
                print(f"Failed to process: {name}, {surname}, {photo_link}, {quotes}")
            
        # if((i-3)%6 == 0):
        #     doc.add_page_break()
              
                   
# 1/0.95
doc.save('facedetec.docx')